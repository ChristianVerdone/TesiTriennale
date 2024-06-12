from docx import Document


def read_docx(file_path):
    doc = Document(file_path)
    listLines = []
    print(doc.paragraphs.__len__())
    for para in doc.paragraphs:
        # Remove unwanted characters and words
        para.text = para.text.replace('!', ' ').replace('  ', ' ').replace('  ', ' ')
        if (para.text.__contains__('SCHEDA') or para.text.__contains__('PAGINA') or
                para.text.__contains__('Codice conto: 7.07.001') or
                para.text.__contains__('---') or para.text.__contains__('Nota') or
                para.text.__contains__('Saldi operazioni del periodo') or
                para.text.__contains__('Saldi al') or
                para.text.__contains__('DATA COD DESCRIZ.') or
                para.text.__contains__('OPERAZ CAU') or
                para.text.__contains__('+') or para.text.__contains__('    507002') or
                para.text.__contains__('     ') or para.text.__contains__('    505013')):
            para.text = ''

    i = 0

    # Remove paragraphs until the first 'Codice' is found
    while not doc.paragraphs[i].text.__contains__('Codice'):
        doc.paragraphs[i].text = ''
        i += 1

    new_doc = Document()  # Create a new Document
    # Add non-empty paragraphs to the new document
    for para in doc.paragraphs:
        if para.text != '':
            new_doc.add_paragraph(para.text)

    i = 0
    for para in new_doc.paragraphs:
        print(str(i) + ' :' + para.text)
        i += 1
    print(new_doc.paragraphs.__len__())
    return listLines


items = read_docx('2022.docx')
